import io

from app.rag import parsers


def test_txt_utf8():
    assert parsers.extract("금요일 배포 금지".encode("utf-8"), "a.txt") == "금요일 배포 금지"


def test_txt_cp949_fallback():
    assert parsers.extract("보류 사유".encode("cp949"), "a.txt") == "보류 사유"


def test_csv_flattened_to_labelled_rows():
    content = "월,생성,완료\n2026-06,58,40\n2026-07,49,42\n".encode("utf-8")
    text = parsers.extract(content, "stats.csv")
    assert "월: 2026-06, 생성: 58, 완료: 40" in text
    # 헤더가 청크마다 반복되지 않아야 검색이 헤더에만 걸리지 않는다.
    assert text.count("월: ") == 2


def test_docx_paragraphs_and_tables():
    import docx

    document = docx.Document()
    document.add_paragraph("배포 절차서 v2")
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "금지"
    table.rows[0].cells[1].text = "금요일 15시"
    buffer = io.BytesIO()
    document.save(buffer)

    text = parsers.extract(buffer.getvalue(), "doc.docx")
    assert "배포 절차서 v2" in text
    assert "금지 | 금요일 15시" in text


def test_xlsx_sheets():
    from openpyxl import Workbook

    workbook = Workbook()
    workbook.active.title = "통계"
    workbook.active.append(["월", "지연"])
    workbook.active.append(["2026-06", 14])
    buffer = io.BytesIO()
    workbook.save(buffer)

    text = parsers.extract(buffer.getvalue(), "s.xlsx")
    assert "[시트: 통계]" in text
    assert "2026-06 | 14" in text


def test_unsupported_extension_rejected():
    assert not parsers.can_extract("image.png")
    try:
        parsers.extract(b"", "image.png")
    except parsers.UnsupportedFormat:
        return
    raise AssertionError("UnsupportedFormat 이 나야 한다")
